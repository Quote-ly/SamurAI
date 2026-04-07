"""File handling — parse inbound documents and create edited versions."""

import io
import logging
import os
import tempfile

from langchain_core.tools import tool

logger = logging.getLogger(__name__)

# Store uploaded file content per conversation for editing
# Key: conversation_id, Value: {filename, content_bytes, file_type, text_content}
_uploaded_files: dict[str, dict] = {}

# Store pending edited files for FileConsentCard upload
# Key: conversation_id, Value: {filename, content_bytes, summary}
_pending_edited_files: dict[str, dict] = {}


# ── File Parsing ──────────────────────────────────────────────────────


def parse_docx(content_bytes: bytes) -> str:
    """Extract text from a .docx file."""
    from docx import Document

    doc = Document(io.BytesIO(content_bytes))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Also extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            paragraphs.append(" | ".join(cells))

    return "\n".join(paragraphs)


def parse_xlsx(content_bytes: bytes) -> str:
    """Extract text from an .xlsx file."""
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(content_bytes), read_only=True, data_only=True)
    lines = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        lines.append(f"**Sheet: {sheet_name}**")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            if any(cells):
                lines.append(" | ".join(cells))
        lines.append("")

    wb.close()
    return "\n".join(lines)


def parse_pdf(content_bytes: bytes) -> str:
    """Extract text from a PDF file."""
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(content_bytes))
    pages = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text and text.strip():
            pages.append(f"[Page {i + 1}]\n{text}")

    return "\n\n".join(pages) if pages else "Could not extract text from PDF."


def parse_file(filename: str, content_bytes: bytes) -> tuple[str, str]:
    """Parse a file and return (text_content, file_type)."""
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".docx":
        return parse_docx(content_bytes), "docx"
    elif ext == ".xlsx":
        return parse_xlsx(content_bytes), "xlsx"
    elif ext == ".pdf":
        return parse_pdf(content_bytes), "pdf"
    elif ext in (".txt", ".md", ".csv", ".json", ".py", ".yaml", ".yml"):
        return content_bytes.decode("utf-8", errors="replace"), ext.lstrip(".")
    else:
        return f"Unsupported file type: {ext}", "unknown"


# ── File Editing ──────────────────────────────────────────────────────


def create_edited_docx(original_bytes: bytes, instructions: str, new_content: str) -> bytes:
    """Create an edited .docx file by replacing content."""
    from docx import Document

    doc = Document(io.BytesIO(original_bytes))

    # Clear existing paragraphs and replace with new content
    # Preserve formatting from the first paragraph style
    first_style = doc.paragraphs[0].style if doc.paragraphs else None

    for para in doc.paragraphs:
        para.clear()

    # Add new content as paragraphs
    for i, line in enumerate(new_content.split("\n")):
        if i < len(doc.paragraphs):
            doc.paragraphs[i].text = line
            if first_style:
                doc.paragraphs[i].style = first_style
        else:
            doc.add_paragraph(line)

    output = io.BytesIO()
    doc.save(output)
    return output.getvalue()


def create_edited_xlsx(original_bytes: bytes, sheet_name: str, updates: list[dict]) -> bytes:
    """Apply updates to an .xlsx file.

    updates: list of {row, col, value} dicts (1-indexed)
    """
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(original_bytes))
    ws = wb[sheet_name] if sheet_name in wb.sheetnames else wb.active

    for update in updates:
        row = update.get("row", 1)
        col = update.get("col", 1)
        value = update.get("value", "")
        ws.cell(row=row, column=col, value=value)

    output = io.BytesIO()
    wb.save(output)
    return output.getvalue()


# ── Agent Tools ───────────────────────────────────────────────────────


@tool
def get_uploaded_file_content(conversation_id: str) -> str:
    """Get the text content of a file uploaded by the user in this conversation.

    Args:
        conversation_id: The current conversation ID (from context brackets).
    """
    file_info = _uploaded_files.get(conversation_id)
    if not file_info:
        return "No file has been uploaded in this conversation."

    filename = file_info["filename"]
    text = file_info["text_content"]
    file_type = file_info["file_type"]

    if len(text) > 30000:
        text = text[:30000] + f"\n\n... [truncated at 30,000 chars, full file is {len(text)} chars]"

    return f"**File: {filename}** (type: {file_type})\n\n{text}"


@tool
def read_spreadsheet_cells(
    conversation_id: str,
    sheet_name: str,
    column: str,
    start_row: int = 2,
    end_row: int = 0,
) -> str:
    """Read specific cells from the uploaded (or edited) spreadsheet.

    Use this to verify edits were applied or to read specific column values.
    Reads from the most recent version (edited if available, otherwise original).

    Args:
        conversation_id: The current conversation ID (from context brackets).
        sheet_name: The sheet to read from.
        column: The header name of the column to read.
        start_row: First data row to read (default 2, headers on row 1).
        end_row: Last row to read (default 0 = all rows).
    """
    from openpyxl import load_workbook

    # Use edited version if available, otherwise original
    file_info = _uploaded_files.get(conversation_id)
    if not file_info:
        return "No file has been uploaded in this conversation."
    if file_info["file_type"] != "xlsx":
        return "This tool only works with .xlsx files."

    try:
        wb = load_workbook(io.BytesIO(file_info["content_bytes"]), read_only=True, data_only=True)
        ws = wb[sheet_name] if sheet_name in wb.sheetnames else wb.active

        # Find column index
        headers = [str(c.value) if c.value else "" for c in ws[1]]
        col_idx = None
        for idx, h in enumerate(headers):
            if h.strip().lower() == column.strip().lower():
                col_idx = idx
                break

        if col_idx is None:
            wb.close()
            return f"Column '{column}' not found. Available: {', '.join(headers)}"

        max_row = ws.max_row
        actual_end = end_row if end_row > 0 else max_row

        lines = [f"**{column}** (rows {start_row}-{actual_end}):\n"]
        for row_num in range(start_row, actual_end + 1):
            val = ws.cell(row=row_num, column=col_idx + 1).value
            val_str = str(val)[:150] if val is not None else "(empty)"
            lines.append(f"Row {row_num}: {val_str}")

        wb.close()
        return "\n".join(lines)
    except Exception as e:
        return f"Error reading cells: {e}"


@tool
def edit_document(
    conversation_id: str,
    user_email: str,
    new_content: str,
    summary: str = "Document edited by SamurAI",
) -> str:
    """Edit the uploaded document and send the modified version back.

    Only works with .docx files. Creates a new version with the provided content
    and sends it back to the user via Teams file upload.

    Args:
        conversation_id: The current conversation ID (from context brackets).
        user_email: The user's email (from context brackets).
        new_content: The full new content for the document.
        summary: Brief description of the changes made.
    """
    file_info = _uploaded_files.get(conversation_id)
    if not file_info:
        return "No file has been uploaded in this conversation to edit."

    file_type = file_info["file_type"]
    if file_type not in ("docx", "xlsx"):
        return f"Cannot edit {file_type} files. Only .docx and .xlsx files can be edited."

    try:
        if file_type == "docx":
            edited_bytes = create_edited_docx(
                file_info["content_bytes"], summary, new_content
            )
        else:
            return "For .xlsx edits, use edit_spreadsheet instead."

        filename = file_info["filename"]
        edited_name = f"edited_{filename}"

        # Store for FileConsentCard upload
        _pending_edited_files[conversation_id] = {
            "filename": edited_name,
            "content_bytes": edited_bytes,
            "summary": summary,
        }

        return f"Document edited. The modified file '{edited_name}' will be sent to you for review."
    except Exception as e:
        return f"Error editing document: {e}"


@tool
def edit_spreadsheet(
    conversation_id: str,
    user_email: str,
    sheet_name: str,
    updates: str,
    summary: str = "Spreadsheet edited by SamurAI",
) -> str:
    """Edit the uploaded spreadsheet and send the modified version back.

    Only works with .xlsx files.

    Args:
        conversation_id: The current conversation ID (from context brackets).
        user_email: The user's email (from context brackets).
        sheet_name: The sheet to edit (use the sheet name from get_uploaded_file_content).
        updates: JSON string of updates as a list of objects with row, col, value.
                 Example: '[{"row": 2, "col": 3, "value": "new value"}]'
                 Rows and columns are 1-indexed.
        summary: Brief description of the changes made.
    """
    import json as _json

    file_info = _uploaded_files.get(conversation_id)
    if not file_info:
        return "No file has been uploaded in this conversation to edit."

    if file_info["file_type"] != "xlsx":
        return "This tool only works with .xlsx files."

    try:
        update_list = _json.loads(updates)
        edited_bytes = create_edited_xlsx(
            file_info["content_bytes"], sheet_name, update_list
        )

        filename = file_info["filename"]
        edited_name = f"edited_{filename}"

        _pending_edited_files[conversation_id] = {
            "filename": edited_name,
            "content_bytes": edited_bytes,
            "summary": summary,
        }

        # Update stored file so subsequent reads/edits see the changes
        file_info["content_bytes"] = edited_bytes
        file_info["text_content"] = parse_xlsx(edited_bytes)

        return (
            f"Spreadsheet edited ({len(update_list)} cells updated). "
            f"The modified file '{edited_name}' will be sent to you.\n"
            "Use read_spreadsheet_cells to verify the changes."
        )
    except Exception as e:
        return f"Error editing spreadsheet: {e}"


@tool
def get_spreadsheet_info(conversation_id: str) -> str:
    """Get the structure of an uploaded spreadsheet — sheets, columns, row count, and sample data.

    Use this before fill_spreadsheet_column to understand the sheet layout.

    Args:
        conversation_id: The current conversation ID (from context brackets).
    """
    file_info = _uploaded_files.get(conversation_id)
    if not file_info:
        return "No file has been uploaded in this conversation."
    if file_info["file_type"] != "xlsx":
        return "This tool only works with .xlsx files."

    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(file_info["content_bytes"]), read_only=True, data_only=True)
    lines = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = list(ws.iter_rows(values_only=True))
        if not rows:
            lines.append(f"**Sheet: {sheet_name}** — empty")
            continue

        headers = [str(c) if c is not None else "" for c in rows[0]]
        row_count = len(rows) - 1  # Exclude header

        lines.append(f"**Sheet: {sheet_name}** — {row_count} data rows, {len(headers)} columns")
        lines.append(f"Columns: {' | '.join(headers)}")

        # Show first 3 data rows as sample
        sample_rows = rows[1:4]
        if sample_rows:
            lines.append("Sample data:")
            for r in sample_rows:
                cells = [str(c) if c is not None else "" for c in r]
                lines.append(f"  {' | '.join(cells)}")

        # Show which columns have empty cells
        empty_cols = []
        for col_idx, header in enumerate(headers):
            empty_count = sum(1 for r in rows[1:] if r[col_idx] is None or str(r[col_idx]).strip() == "")
            if empty_count > 0:
                empty_cols.append(f"{header} ({empty_count} empty)")
        if empty_cols:
            lines.append(f"Columns with empty cells: {', '.join(empty_cols)}")
        lines.append("")

    wb.close()
    return "\n".join(lines)


@tool
def fill_spreadsheet_column(
    conversation_id: str,
    user_email: str,
    sheet_name: str,
    target_column: str,
    expression: str,
    start_row: int = 2,
    end_row: int = 0,
    summary: str = "Column filled by SamurAI",
) -> str:
    """Fill a column in the uploaded spreadsheet using a Python expression.

    The expression is evaluated for each row with access to all column values.
    Column values are available by their header name (spaces replaced with underscores,
    lowercased). Example: if headers are "Name", "Score", "Grade", the expression
    can reference `name`, `score`, `grade`.

    Built-in helpers available in the expression:
    - row_num: the current row number (1-indexed)
    - All standard Python: str(), int(), float(), len(), round(), etc.

    Examples:
        target_column="Grade", expression="'A' if score >= 90 else 'B' if score >= 80 else 'C'"
        target_column="Full Name", expression="f'{first_name} {last_name}'"
        target_column="Risk", expression="'High' if severity > 7 else 'Medium' if severity > 4 else 'Low'"
        target_column="Status", expression="'Overdue' if days_open > 30 else 'On Track'"

    Args:
        conversation_id: The current conversation ID (from context brackets).
        user_email: The user's email (from context brackets).
        sheet_name: The sheet to edit.
        target_column: The header name of the column to fill.
        expression: A Python expression evaluated per row. Column values available as variables.
        start_row: First data row to fill (default 2, since row 1 is headers).
        end_row: Last row to fill (default 0 = all rows).
        summary: Brief description of the changes.
    """
    from openpyxl import load_workbook

    file_info = _uploaded_files.get(conversation_id)
    if not file_info:
        return "No file has been uploaded in this conversation."
    if file_info["file_type"] != "xlsx":
        return "This tool only works with .xlsx files."

    try:
        wb = load_workbook(io.BytesIO(file_info["content_bytes"]))
        ws = wb[sheet_name] if sheet_name in wb.sheetnames else wb.active

        # Get headers from row 1
        headers = []
        for cell in ws[1]:
            val = str(cell.value) if cell.value is not None else ""
            headers.append(val)

        # Map header names to column indices
        header_map = {}
        target_col_idx = None
        for idx, h in enumerate(headers):
            safe_name = h.strip().lower().replace(" ", "_").replace("-", "_")
            safe_name = "".join(c for c in safe_name if c.isalnum() or c == "_")
            header_map[safe_name] = idx
            if h.strip().lower() == target_column.strip().lower():
                target_col_idx = idx

        if target_col_idx is None:
            return f"Column '{target_column}' not found. Available columns: {', '.join(headers)}"

        # Determine row range
        max_row = ws.max_row
        actual_end = end_row if end_row > 0 else max_row

        filled = 0
        errors = []

        for row_num in range(start_row, actual_end + 1):
            # Build context variables from this row's cells
            row_vars = {"row_num": row_num}
            for safe_name, col_idx in header_map.items():
                cell_val = ws.cell(row=row_num, column=col_idx + 1).value
                row_vars[safe_name] = cell_val if cell_val is not None else ""

            try:
                # Evaluate the expression safely
                result = eval(expression, {"__builtins__": {
                    "str": str, "int": int, "float": float, "bool": bool,
                    "len": len, "round": round, "abs": abs, "min": min, "max": max,
                    "sum": sum, "sorted": sorted, "enumerate": enumerate,
                    "True": True, "False": False, "None": None,
                }}, row_vars)

                ws.cell(row=row_num, column=target_col_idx + 1, value=result)
                filled += 1
                if filled <= 3:
                    print(f"[fill_spreadsheet] row={row_num} col={target_col_idx + 1} value={str(result)[:80]}", flush=True)
            except Exception as e:
                if len(errors) < 3:
                    errors.append(f"Row {row_num}: {e}")

        # Save
        output = io.BytesIO()
        wb.save(output)
        edited_bytes = output.getvalue()
        wb.close()

        filename = file_info["filename"]
        edited_name = f"edited_{filename}"

        _pending_edited_files[conversation_id] = {
            "filename": edited_name,
            "content_bytes": edited_bytes,
            "summary": summary,
        }

        # Update the stored file so subsequent reads/edits see the changes
        file_info["content_bytes"] = edited_bytes
        file_info["text_content"] = parse_xlsx(edited_bytes)

        result_msg = f"Filled {filled} cells in column '{target_column}'."
        if errors:
            result_msg += f"\n{len(errors)} errors: " + "; ".join(errors)
        result_msg += f"\nThe modified file '{edited_name}' will be sent to you."
        result_msg += "\nUse read_spreadsheet_cells to verify the changes."
        return result_msg

    except Exception as e:
        return f"Error: {e}"


FILE_HANDLER_TOOLS = [
    get_uploaded_file_content,
    get_spreadsheet_info,
    read_spreadsheet_cells,
    edit_document,
    edit_spreadsheet,
    fill_spreadsheet_column,
]
